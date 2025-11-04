import fitz  # PyMuPDF
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document
import re
from typing import Dict, Any

class TextExtractor:
    """Extract text from PDF and DOCX files"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF (fallback to pdfminer)"""
        try:
            # Try PyMuPDF first
            doc = fitz.open(file_path)
            text = ""
            metadata = {
                "pages": len(doc),
                "has_images": False,
                "has_tables": False
            }
            
            for page_num, page in enumerate(doc):
                text += page.get_text()
                
                # Check for images
                if page.get_images():
                    metadata["has_images"] = True
            
            doc.close()
            
            # If PyMuPDF fails, use pdfminer
            if not text.strip():
                text = pdf_extract_text(file_path)
            
            return {
                "text": text.strip(),
                "metadata": metadata,
                "extraction_method": "PyMuPDF"
            }
        except Exception as e:
            # Fallback to pdfminer
            try:
                text = pdf_extract_text(file_path)
                return {
                    "text": text.strip(),
                    "metadata": {"pages": 0},
                    "extraction_method": "pdfminer"
                }
            except Exception as fallback_error:
                raise Exception(f"Failed to extract PDF: {str(e)}, Fallback: {str(fallback_error)}")
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "has_images": any(rel.target_ref for rel in doc.part.rels.values() 
                                if "image" in rel.target_ref)
            }
            
            return {
                "text": text.strip(),
                "metadata": metadata,
                "extraction_method": "python-docx"
            }
        except Exception as e:
            raise Exception(f"Failed to extract DOCX: {str(e)}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean extracted text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\-\@\+\(\)\/\:]', '', text)
        return text.strip()
    
    def extract(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Main extraction method"""
        if file_type == 'pdf':
            result = self.extract_from_pdf(file_path)
        elif file_type in ['doc', 'docx']:
            result = self.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Clean the extracted text
        result['text'] = self.clean_text(result['text'])
        result['word_count'] = len(result['text'].split())
        
        return result
